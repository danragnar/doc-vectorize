#!env python
import os
import sys
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
import torch

from langchain.chains import RetrievalQA
import faiss

def load_pdf(file_path):
    """Extract text from PDF file."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def load_docx(file_path):
    """Extract text from DOCX file."""
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    return text

def extract_doc_metadata(file_path):
    """Extract creation and modification dates from document metadata."""
    try:
        if file_path.endswith('.pdf'):
            reader = PdfReader(file_path)
            info = reader.metadata
            created = getattr(info, 'creation_date', None)
            modified = getattr(info, 'modification_date', None)
            print(created)
            return {
                "created": created.isoformat() if created else None,
                "modified": modified.isoformat() if modified else None
            }
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            props = doc.core_properties
            return {
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None
            }
    except Exception:
        pass
    return {}

def load_document(file_path):
    """Load document based on file extension."""
    if file_path.endswith('.pdf'):
        return load_pdf(file_path)
    elif file_path.endswith('.docx'):
        return load_docx(file_path)
    else:
        raise ValueError("Unsupported file type")

def chunk_text(text, chunk_size=1000, chunk_overlap=200):
    """Split text into chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return splitter.split_text(text)

def generate_embeddings_local(chunks, model_name='Qwen/Qwen3-Embedding-0.6B'):
    """Generate embeddings using local sentence-transformers model."""
    model = SentenceTransformer('Qwen/Qwen3-Embedding-0.6B', model_kwargs={"attn_implementation": "flash_attention_2", "device_map": "auto", "torch_dtype": torch.float16}, tokenizer_kwargs={"padding_side": "left"} )
    embeddings = model.encode(chunks)
    return embeddings

def create_vector_store(chunks, embeddings):
    """Create FAISS vector store."""
    if embeddings.size == 0:
        raise ValueError("No embeddings to create vector store")
    # Create FAISS index
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    # Store chunks as metadata
    return index, chunks

def save_vector_store(index, metadata, save_path='vector_store.index'):
    """Save FAISS index and metadata."""
    faiss.write_index(index, save_path)
    import json
    with open(save_path + '.metadata', 'w') as f:
        json.dump(metadata, f)

def load_vector_store(load_path='vector_store.index'):
    """Load FAISS index and metadata."""
    index = faiss.read_index(load_path)
    import json
    with open(load_path + '.metadata', 'r') as f:
        metadata = json.load(f)
    return index, metadata

def search_similar(query_embedding, index, metadata, k=5):
    """Search for similar metadata."""
    query_emb = query_embedding.reshape(1, -1)
    distances, indices = index.search(query_emb, k)
    results = [metadata[i] for i in indices[0]]
    return results

def query_pipeline(query, index, metadata, model_name='Qwen/Qwen3-Embedding-0.6B', openai_api_key=None):
    """Query pipeline with retrieval and ChatGPT."""
    import openai
    
    # Generate query embedding
    model = SentenceTransformer('Qwen/Qwen3-Embedding-0.6B', model_kwargs={"attn_implementation": "flash_attention_2", "device_map": "auto", "torch_dtype": torch.float16}, tokenizer_kwargs={"padding_side": "left"} )
    query_emb = model.encode([query])[0]
    
    # Retrieve similar metadata
    relevant_metadata = search_similar(query_emb, index, metadata)
    context_parts = []
    for item in relevant_metadata:
        file_info = f"From {item['file']}"
        if item.get('created'):
            file_info += f" (created: {item['created']})"
        if item.get('modified'):
            file_info += f" (modified: {item['modified']})"
        context_parts.append(f"{file_info}: {item['text']}")
    context = "\n".join(context_parts)
    
    # Use ChatGPT for answer
    if not openai_api_key:
        openai_api_key = os.getenv('OPENAI_API_KEY')
    if not openai_api_key:
        raise ValueError("OpenAI API key required")
    
    openai.api_key = openai_api_key
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": f"Context: {context}\n\nQuestion: {query}\n\nAnswer:"}]
    )
    return response.choices[0].message.content

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Document Vectorizer")
    subparsers = parser.add_subparsers(dest='command')

    # Ingest
    ingest_parser = subparsers.add_parser('ingest')
    ingest_parser.add_argument('input_dir')

    # Query
    query_parser = subparsers.add_parser('query')
    query_parser.add_argument('question', nargs='+')

    # Search
    search_parser = subparsers.add_parser('search')
    search_parser.add_argument('question', nargs='+')
    search_parser.add_argument('--from-date', help='From date (YYYY-MM-DD)')
    search_parser.add_argument('--to-date', help='To date (YYYY-MM-DD)')

    # Chat
    chat_parser = subparsers.add_parser('chat')

    args = parser.parse_args()
    if not args.command:
        parser.print_help()
        return

    command = args.command
    
    if command == 'ingest':
        input_dir = args.input_dir
        all_metadata = []
        all_embeddings = []
        
        # Create stored_documents folder
        stored_dir = 'stored_documents'
        os.makedirs(stored_dir, exist_ok=True)
        
        import hashlib
        import shutil
        
        for file in os.listdir(input_dir):
            if file.endswith(('.pdf', '.docx')):
                file_path = os.path.join(input_dir, file)
                
                # Compute SHA256 hash
                with open(file_path, 'rb') as f:
                    file_hash = hashlib.sha256(f.read()).hexdigest()
                
                ext = file.split('.')[-1]
                subdir = file_hash[:2]
                subdir_path = os.path.join(stored_dir, subdir)
                os.makedirs(subdir_path, exist_ok=True)
                stored_filename = f"{file_hash}.{ext}"
                stored_path = os.path.join(subdir_path, stored_filename)
                
                # Copy if not exists
                if not os.path.exists(stored_path):
                    shutil.copy2(file_path, stored_path)
                
                text = load_document(file_path)
                doc_metadata = extract_doc_metadata(file_path)
                chunks = chunk_text(text)
                embeddings = generate_embeddings_local(chunks)
                
                for chunk in chunks:
                    all_metadata.append({"text": chunk, "file": file, "stored_path": stored_path, **doc_metadata})
                all_embeddings.extend(embeddings)
        
        # Convert to numpy array
        import numpy as np
        all_embeddings = np.array(all_embeddings)
        index, metadata = create_vector_store(all_metadata, all_embeddings)
        save_vector_store(index, metadata)
        print("Ingestion complete")
    
    elif command == 'query':
        question = ' '.join(args.question)
        index, metadata = load_vector_store()
        answer = query_pipeline(question, index, metadata)
        print(answer)
    
    elif command == 'search':
        question = ' '.join(args.question)
        from_date = args.from_date
        to_date = args.to_date
        index, metadata = load_vector_store()
        model = SentenceTransformer('Qwen/Qwen3-Embedding-0.6B', model_kwargs={"attn_implementation": "flash_attention_2", "device_map": "auto", "torch_dtype": torch.float16}, tokenizer_kwargs={"padding_side": "left"} )
        query_emb = model.encode([question])[0]
        relevant_metadata = search_similar(query_emb, index, metadata)
        
        # Filter by date
        if from_date or to_date:
            from datetime import datetime
            filtered = []
            for item in relevant_metadata:
                modified = item.get('modified')
                if modified:
                    try:
                        mod_dt = datetime.fromisoformat(modified.replace('Z', '+00:00'))
                        if from_date:
                            from_dt = datetime.fromisoformat(from_date)
                            if mod_dt < from_dt:
                                continue
                        if to_date:
                            to_dt = datetime.fromisoformat(to_date)
                            if mod_dt > to_dt:
                                continue
                        filtered.append(item)
                    except:
                        continue  # Skip if date parsing fails
                else:
                    filtered.append(item)  # Include if no date
            relevant_metadata = filtered
        
        print("Relevant chunks:")
        for i, item in enumerate(relevant_metadata, 1):
            date_info = f" (modified: {item.get('modified', 'unknown')})" if item.get('modified') else ""
            print(f"{i}. {item['text']} (from {item['file']}){date_info}")
            print("---")
    
    elif command == 'chat':
        index, metadata = load_vector_store()
        history = []
        print("Start chatting! Type 'exit' or 'quit' to end.")
        while True:
            try:
                question = input("You: ")
            except EOFError:
                break
            if question.lower() in ['exit', 'quit']:
                break
            
            # Retrieve context
            model = SentenceTransformer('Qwen/Qwen3-Embedding-0.6B', model_kwargs={"attn_implementation": "flash_attention_2", "device_map": "auto", "torch_dtype": torch.float16}, tokenizer_kwargs={"padding_side": "left"} )
            query_emb = model.encode([question])[0]
            relevant_metadata = search_similar(query_emb, index, metadata)
            context_parts = []
            for item in relevant_metadata:
                file_info = f"From {item['file']}"
                if item.get('created'):
                    file_info += f" (created: {item['created']})"
                if item.get('modified'):
                    file_info += f" (modified: {item['modified']})"
                context_parts.append(f"{file_info}: {item['text']}")
            context = "\n".join(context_parts)
            
            # Build messages with history
            messages = [{"role": "system", "content": f"Context from documents: {context}"}] + history + [
                {"role": "user", "content": question}
            ]
            
            # Call OpenAI
            import openai
            openai.api_key = os.getenv('OPENAI_API_KEY')
            if not openai.api_key:
                print("Error: OPENAI_API_KEY not set.")
                break
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            answer = response.choices[0].message.content
            
            print(f"AI: {answer}")
            history.append({"role": "user", "content": question})
            history.append({"role": "assistant", "content": answer})

if __name__ == "__main__":
    main()
